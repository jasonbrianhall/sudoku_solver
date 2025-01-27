from sudoku_solver import Sudoku, PuzzleGenerator
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from typing import Dict

class SudokuPuzzleGenerator:
    VALID_DIFFICULTIES = ['easy', 'medium', 'hard', 'extreme']
    
    def __init__(self):
        self.solver = Sudoku()
        self.generator = PuzzleGenerator(self.solver)
        
    def generate_puzzle(self, difficulty="easy"):
        """Generate a new Sudoku puzzle with given difficulty"""
        if difficulty not in self.VALID_DIFFICULTIES:
            raise ValueError(f"Invalid difficulty. Must be one of {self.VALID_DIFFICULTIES}")
            
        # Generate the puzzle
        if not self.generator.generate_puzzle(difficulty):
            raise Exception("Failed to generate puzzle with difficulty: " + difficulty)
            
        # Get unsolved puzzle state
        puzzle = self._get_current_grid()
        
        # Solve to get solution
        if self.solver.solve() != 0:
            raise Exception("Failed to solve puzzle")
            
        # Get solution grid
        solution = self._get_current_grid()
        
        return puzzle, solution
    
    def _get_current_grid(self):
        return [[self.solver.get_value(i, j) for j in range(9)] for i in range(9)]

    def _set_cell_border(self, cell, top=None, right=None, bottom=None, left=None):
        """Helper function to set cell border properties"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # Create dictionaries for border styles
        thin_border = {'sz': '4', 'val': 'single', 'color': '000000'}
        thick_border = {'sz': '40', 'val': 'single', 'color': '000000'}
        
        def create_border_element(edge, attrs):
            return parse_xml(f'<w:{edge} {nsdecls("w")} w:sz="{attrs["sz"]}" w:val="{attrs["val"]}" w:color="{attrs["color"]}"/>')
        
        # Set each border
        if top:
            tcPr.append(create_border_element('top', thick_border if top == 'thick' else thin_border))
        if right:
            tcPr.append(create_border_element('right', thick_border if right == 'thick' else thin_border))
        if bottom:
            tcPr.append(create_border_element('bottom', thick_border if bottom == 'thick' else thin_border))
        if left:
            tcPr.append(create_border_element('left', thick_border if left == 'thick' else thin_border))

    def _format_table(self, table):
        """Apply formatting to make the Sudoku grid look proper"""
        # Set table properties
        table.allow_autofit = False
        table.width = Inches(4.5)  # Make table square
        
        # Set consistent cell size and center alignment
        for row in table.rows:
            row.height = Inches(0.5)  # Make cells square
            for cell in row.cells:
                cell.width = Inches(0.5)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.font.size = Pt(14)
                run.font.name = 'Arial'

        # Apply borders
        for i in range(9):
            for j in range(9):
                cell = table.cell(i, j)
                
                # Determine border thickness for each side
                top = 'thick' if i == 0 or i % 3 == 0 else 'thin'
                bottom = 'thick' if i == 8 or (i + 1) % 3 == 0 else 'thin'
                left = 'thick' if j == 0 or j % 3 == 0 else 'thin'
                right = 'thick' if j == 8 or (j + 1) % 3 == 0 else 'thin'
                
                # Apply borders to cell
                self._set_cell_border(cell, top=top, right=right, bottom=bottom, left=left)

    def create_word_document(self, puzzle_counts: Dict[str, int], filename="sudoku_puzzles.docx"):
        """Create a Word document with specified number of puzzles for each difficulty level
        
        Args:
            puzzle_counts: Dictionary with difficulty levels as keys and number of puzzles as values
                         e.g., {'easy': 2, 'medium': 3, 'hard': 1, 'extreme': 1}
            filename: Output filename for the Word document
        """
        doc = Document()
        doc.add_heading('Sudoku Puzzles', 0)
        
        # Validate difficulties
        for difficulty in puzzle_counts.keys():
            if difficulty not in self.VALID_DIFFICULTIES:
                raise ValueError(f"Invalid difficulty '{difficulty}'. Must be one of {self.VALID_DIFFICULTIES}")
        
        # Store puzzles and solutions for later
        all_puzzles = []  # List of tuples: (difficulty, puzzle_number, puzzle_grid)
        all_solutions = []  # List of tuples: (difficulty, puzzle_number, solution_grid)
        
        # Generate all puzzles first
        for difficulty, count in puzzle_counts.items():
            for i in range(count):
                puzzle, solution = self.generate_puzzle(difficulty)
                puzzle_num = len(all_puzzles) + 1
                all_puzzles.append((difficulty, puzzle_num, puzzle))
                all_solutions.append((difficulty, puzzle_num, solution))
        
        # Add all puzzles
        doc.add_heading('Puzzles', 1)
        for difficulty, puzzle_num, puzzle in all_puzzles:
            doc.add_heading(f'Puzzle {puzzle_num} ({difficulty})', 2)
            
            table = doc.add_table(rows=9, cols=9)
            table.style = 'Table Grid'
            
            for row in range(9):
                for col in range(9):
                    cell = table.cell(row, col)
                    value = puzzle[row][col]
                    cell.text = str(value + 1) if value >= 0 else ''
            
            self._format_table(table)
            
            # Add page break after each puzzle except the last one
            if puzzle_num < len(all_puzzles):
                doc.add_page_break()
        
        # Add page break before solutions section
        doc.add_page_break()
        
        # Add all solutions
        doc.add_heading('Solutions', 1)
        for difficulty, puzzle_num, solution in all_solutions:
            doc.add_heading(f'Solution {puzzle_num} ({difficulty})', 2)
            
            table = doc.add_table(rows=9, cols=9)
            table.style = 'Table Grid'
            
            for row in range(9):
                for col in range(9):
                    cell = table.cell(row, col)
                    value = solution[row][col]
                    cell.text = str(value + 1) if value >= 0 else ''
            
            self._format_table(table)
            
            # Add page break after each solution except the last one
            if puzzle_num < len(all_solutions):
                doc.add_page_break()
        
        doc.save(filename)
